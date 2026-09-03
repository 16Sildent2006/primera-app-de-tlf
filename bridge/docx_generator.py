import os
import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from bridge.oriente_engine import TripQuote

def set_cell_background(cell, fill_hex):
    """Establece el color de fondo de una celda en Word."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def generate_quote_docx(quote: TripQuote, params: dict, output_dir: str = "cotizaciones") -> str:
    """
    Genera un documento Word (.docx) profesional con toda la cotización y ruta.
    Retorna la ruta absoluta del archivo generado.
    """
    os.makedirs(output_dir, exist_ok=True)

    timestamp_str = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    date_display = datetime.datetime.now().strftime("%d/%m/%Y %I:%M %p")
    quote_id = f"COT-350-{timestamp_str[-6:]}"

    orig_clean = quote.origin.name.replace(" ", "_").replace("í", "i").replace("á", "a").replace("é", "e").replace("ó", "o").replace("ú", "u").replace("ü", "u")
    dest_clean = quote.destination.name.replace(" ", "_").replace("í", "i").replace("á", "a").replace("é", "e").replace("ó", "o").replace("ú", "u").replace("ü", "u")
    filename = f"Cotizacion_350_{orig_clean}_a_{dest_clean}_{timestamp_str}.docx"
    file_path = os.path.abspath(os.path.join(output_dir, filename))

    doc = Document()

    # Ajustar márgenes
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # --- TÍTULO PRINCIPAL ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("FLETES ORIENTE - TRANSPORTE & LOGÍSTICA\n")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(24, 43, 73) # Azul corporativo

    run_sub = title_p.add_run("Comprobante Oficial de Cotización • Camión 350\n")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(12)
    run_sub.font.color.rgb = RGBColor(100, 110, 125)

    # Meta
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_meta = meta_p.add_run(f"N° Cotización: {quote_id}  |  Fecha: {date_display}")
    run_meta.font.name = "Arial"
    run_meta.font.size = Pt(9)
    run_meta.font.italic = True
    run_meta.font.color.rgb = RGBColor(120, 120, 120)

    # --- SECCIÓN 1: DATOS DEL VIAJE ---
    h1 = doc.add_paragraph()
    r_h1 = h1.add_run("1. INFORMACIÓN GENERAL DEL VIAJE")
    r_h1.font.name = "Arial"
    r_h1.font.size = Pt(12)
    r_h1.font.bold = True
    r_h1.font.color.rgb = RGBColor(24, 43, 73)

    t1 = doc.add_table(rows=3, cols=2)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    t1.autofit = True

    modalidad = "Ida y Vuelta (Retorno)" if params.get("is_round_trip", False) else "Solo Ida"
    total_mins = int(quote.estimated_hours * 60)
    hours = total_mins // 60
    mins = total_mins % 60
    time_fmt = f"{hours}h {mins}m aprox."

    t1.rows[0].cells[0].text = f"• Origen: {quote.origin.name} ({quote.origin.state})"
    t1.rows[0].cells[1].text = f"• Destino: {quote.destination.name} ({quote.destination.state})"
    t1.rows[1].cells[0].text = f"• Modalidad: {modalidad}"
    t1.rows[1].cells[1].text = f"• Distancia Total: {quote.effective_distance_km:.1f} km"
    t1.rows[2].cells[0].text = f"• Tiempo Estimado: {time_fmt}"
    t1.rows[2].cells[1].text = f"• Combustible Estimado: {quote.fuel_liters:.1f} Litros"

    for row in t1.rows:
        for cell in row.cells:
            set_cell_background(cell, "F2F5F9")
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(10)

    doc.add_paragraph() # Espacio

    # --- SECCIÓN 2: ESPECIFICACIONES DEL CAMIÓN 350 ---
    h2 = doc.add_paragraph()
    r_h2 = h2.add_run("2. ESPECIFICACIONES DEL VEHÍCULO Y CARGA")
    r_h2.font.name = "Arial"
    r_h2.font.size = Pt(12)
    r_h2.font.bold = True
    r_h2.font.color.rgb = RGBColor(24, 43, 73)

    cargo_pct = params.get("cargo_percent", 100.0)
    kg_approx = int(3500 * (cargo_pct / 100.0))

    t2 = doc.add_table(rows=2, cols=2)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    t2.rows[0].cells[0].text = "• Vehículo: Camión 350 (Batea / Plataforma)"
    t2.rows[0].cells[1].text = "• Capacidad Máxima: ~3.500 kg"
    t2.rows[1].cells[0].text = f"• Nivel de Carga: {int(cargo_pct)}%"
    t2.rows[1].cells[1].text = f"• Peso Estimado Carga: ~{kg_approx} kg"

    for row in t2.rows:
        for cell in row.cells:
            set_cell_background(cell, "F2F5F9")
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(10)

    doc.add_paragraph() # Espacio

    # --- SECCIÓN 3: TABLA DE COSTOS Y COTIZACIÓN ---
    h3 = doc.add_paragraph()
    r_h3 = h3.add_run("3. DESGLOSE ECONÓMICO Y TARIFARIO")
    r_h3.font.name = "Arial"
    r_h3.font.size = Pt(12)
    r_h3.font.bold = True
    r_h3.font.color.rgb = RGBColor(24, 43, 73)

    price_km = params.get("price_per_km", 1.25)
    base_fee = params.get("base_fee", 25.0)
    extras = params.get("extra_expenses", 0.0)

    items = [
        ("Tarifa Base de Movilización / Arranque", f"$ {base_fee:.2f} USD"),
        (f"Costo por Distancia ({quote.distance_km:.1f} km a ${price_km:.2f}/km)", f"$ {(quote.distance_km * price_km):.2f} USD"),
        (f"Recargo por Nivel de Carga / Peso ({int(cargo_pct)}% batea)", f"$ {quote.load_surcharge:.2f} USD"),
    ]

    if params.get("is_round_trip", False):
        items.append(("Costo de Retorno (con 30% descuento viaje vacío)", f"$ {quote.return_cost:.2f} USD"))

    if extras > 0:
        items.append(("Gastos Adicionales (Peajes / Ayudantes / Imprevistos)", f"$ {extras:.2f} USD"))

    t3 = doc.add_table(rows=len(items) + 2, cols=2)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Encabezado
    t3.rows[0].cells[0].text = "Concepto"
    t3.rows[0].cells[1].text = "Monto (USD)"
    set_cell_background(t3.rows[0].cells[0], "182B49")
    set_cell_background(t3.rows[0].cells[1], "182B49")

    for cell in t3.rows[0].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(10)
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    for idx, (concepto, monto) in enumerate(items, start=1):
        row = t3.rows[idx]
        row.cells[0].text = concepto
        row.cells[1].text = monto
        bg = "FFFFFF" if idx % 2 != 0 else "F9FAFB"
        set_cell_background(row.cells[0], bg)
        set_cell_background(row.cells[1], bg)
        for p in row.cells[1].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(9.5)

    # Fila de Total
    total_row = t3.rows[-1]
    total_row.cells[0].text = "PRECIO TOTAL A COBRAR:"
    total_row.cells[1].text = f"$ {quote.total_price:.2f} USD"
    set_cell_background(total_row.cells[0], "15803D") # Verde corporativo
    set_cell_background(total_row.cells[1], "15803D")

    for cell in total_row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(12)
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
    for p in total_row.cells[1].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph() # Espacio

    # --- SECCIÓN 4: ITINERARIO DE RUTA ---
    h4 = doc.add_paragraph()
    r_h4 = h4.add_run("4. ITINERARIO DE RUTA VIAL EN ORIENTE")
    r_h4.font.name = "Arial"
    r_h4.font.size = Pt(12)
    r_h4.font.bold = True
    r_h4.font.color.rgb = RGBColor(24, 43, 73)

    itin_p = doc.add_paragraph()
    for i, node in enumerate(quote.path_nodes):
        is_first = (i == 0)
        is_last = (i == len(quote.path_nodes) - 1)
        prefix = "[ORIGEN] " if is_first else ("[DESTINO] " if is_last else f"[{i+1}] ")
        r = itin_p.add_run(f"  {prefix}{node.name} ({node.state}) - Coordenadas: {node.latitude:.4f}, {node.longitude:.4f}\n")
        r.font.name = "Arial"
        r.font.size = Pt(9.5)
        if is_first or is_last:
            r.font.bold = True

    # Guardar documento
    doc.save(file_path)
    return file_path
